 
"""
Hostel Room Booking System - Professional Project Report Generator
Generates a complete Word document report.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, glob

doc = Document()
BASE = r'c:\Users\Anees\Desktop\OOP_ Project'
SRC  = os.path.join(BASE, 'src')
IMGS = os.path.join(BASE, 'screenshots')

# ── Page setup ────────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

# ── Helpers ───────────────────────────────────────────────────────────────────

def h1(text, color=(30, 64, 175)):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(*color)
    return p

def h2(text, color=(37, 99, 235)):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(*color)
    return p

def body(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(11)
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text[:4000])
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    p._p.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Cm(0.5)

def read_src(rel_path):
    try:
        with open(os.path.join(SRC, rel_path), encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return '// File not found: ' + rel_path

def add_image(filename, caption='', width=Inches(5.8)):
    path = os.path.join(IMGS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size = Pt(9)
                r.font.italic = True
    else:
        doc.add_paragraph(f'[Screenshot: {filename} — not found]')

# ==============================================================================
# TITLE PAGE
# ==============================================================================
doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
ru = tp.add_run('COMSATS University Islamabad')
ru.font.size = Pt(16); ru.font.bold = True; ru.font.color.rgb = RGBColor(30,64,175)

doc.add_paragraph()
tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ru2 = tp2.add_run('HOSTEL ROOM BOOKING SYSTEM')
ru2.font.size = Pt(24); ru2.font.bold = True; ru2.font.color.rgb = RGBColor(15,23,42)

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
ru3 = tp3.add_run('Final Project Report')
ru3.font.size = Pt(14); ru3.font.color.rgb = RGBColor(100,116,139)

doc.add_paragraph()
for k, v in [
    ('Course',     'CSC-241 — Object Oriented Programming'),
    ('Semester',   '3rd Semester'),
    ('Language',   'Java SE 11+ with Swing GUI'),
    ('Persistence','File-Based I/O (Pipe-Delimited Text Files)'),
    ('IDE',        'Any Java IDE or Command Line'),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = p.add_run(k + ':  '); rk.font.bold = True; rk.font.size = Pt(12)
    rv = p.add_run(v);          rv.font.size = Pt(12)

doc.add_page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================
h1('Table of Contents')
toc = [
    '1. Introduction',
    '2. Problem Statement & Objectives',
    '3. OOP Concepts Applied (Lab Mapping)',
    '4. System Architecture',
    '5. Project File Structure',
    '6. Key Source Code',
    '7. System Features',
    '8. How to Compile and Run',
    '9. Sample Data Format',
    '10. Output Screenshots',
    '11. Future Enhancements',
    '12. Conclusion',
]
for t in toc:
    doc.add_paragraph(t, style='List Number')
doc.add_page_break()

# ==============================================================================
# 1. INTRODUCTION
# ==============================================================================
h1('1. Introduction')
body(
    'This report documents the design and implementation of the Hostel Room Booking System, '
    'developed as the final semester project for CSC-241 Object Oriented Programming at '
    'COMSATS University Islamabad.\n\n'
    'The system is a fully functional Java Swing desktop application that automates hostel '
    'room management, student registration, room bookings, maintenance tracking, reviews, '
    'and notifications. It is built entirely using Java SE standard libraries with no '
    'external frameworks or databases required. Data persistence is achieved through '
    'pipe-delimited plain text files, fully demonstrating the Java File I/O concepts '
    'covered in the lab curriculum.\n\n'
    'The application supports two distinct user roles: Administrator and Student, each '
    'with a dedicated dashboard tailored to their specific responsibilities.'
)

# ==============================================================================
# 2. PROBLEM STATEMENT
# ==============================================================================
h1('2. Problem Statement & Objectives')
h2('2.1 Problem Statement')
body(
    'Manual hostel management using paper-based or spreadsheet systems is inefficient, '
    'error-prone, and difficult to scale. Key problems include:\n'
)
for p in [
    'No centralized record of room availability leading to double-bookings.',
    'Students have no self-service portal; must physically visit the office.',
    'Maintenance requests are tracked on paper with no status updates.',
    'No structured review or feedback system for rooms.',
    'No automated notification mechanism for important updates.',
]:
    bullet(p)

h2('2.2 Objectives')
for o in [
    'Implement a dual-role system (Admin and Student) with secure login.',
    'Allow students to self-register, browse available rooms, and make bookings.',
    'Enable admins to manage rooms, view all bookings, and monitor students.',
    'Support maintenance request submission and admin-side resolution tracking.',
    'Provide a room review and rating system (1–5 stars).',
    'Implement an in-app notification system for booking and maintenance events.',
    'Persist all data using Java File I/O — no external database required.',
    'Apply all 13+ OOP lab concepts taught in CSC-241.',
]:
    bullet(o)
doc.add_page_break()

# ==============================================================================
# 3. OOP CONCEPTS
# ==============================================================================
h1('3. OOP Concepts Applied — Lab Mapping')
table = doc.add_table(rows=1, cols=4)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['Lab', 'OOP Concept', 'Where Applied', 'File(s)']):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.font.bold = True

rows = [
    ('Lab 03', 'Encapsulation',     'Private fields with getters/setters',            'AbstractPerson, AbstractRoom, all models'),
    ('Lab 04', 'Constructors',      'Parameterized constructors throughout',           'All model classes'),
    ('Lab 05', 'Composition',       'Booking HAS-A Student and AbstractRoom',          'Booking.java'),
    ('Lab 06', 'Inheritance',       'Student/Admin extend AbstractPerson; room types extend AbstractRoom', 'models/'),
    ('Lab 07', 'Abstract Classes',  'Abstract getRoomType(), getMaxOccupancy(), getRole()', 'AbstractRoom, AbstractPerson'),
    ('Lab 08', 'Polymorphism',      'RoomPrinter processes AbstractRoom list; dynamic dispatch', 'Main.java, utils/'),
    ('Lab 09', 'Interfaces',        'Bookable, Saveable, Maintainable, Reviewable, Notifiable', 'interfaces/'),
    ('Lab 10', 'Exception Handling','6 custom checked exceptions for business rules',  'exceptions/'),
    ('Lab 11', 'Generics',          'DataStore<T> generic container used in managers', 'utils/DataStore.java'),
    ('Lab 12', 'File I/O',          'Pipe-delimited text files for all persistence',   'services/FileManager.java'),
    ('Lab 13', 'GUI Layouts',       'BorderLayout, GridBagLayout, CardLayout, FlowLayout', 'gui/LoginFrame.java'),
    ('Lab 14', 'Event Handling',    'ActionListeners, MouseListeners on all controls', 'All GUI files'),
    ('Bonus',  'Design Pattern',    'Service-Oriented Architecture (Manager layer)',   'services/'),
]
for row in rows:
    r = table.add_row().cells
    for i, v in enumerate(row):
        r[i].text = v

doc.add_page_break()

# ==============================================================================
# 4. ARCHITECTURE
# ==============================================================================
h1('4. System Architecture')
body('The system follows a 5-layer Service-Oriented Architecture:\n')
arch = (
    'LAYER 1 — Models (Data/Entity Layer)\n'
    '  AbstractPerson  →  Admin, Student\n'
    '  AbstractRoom    →  SingleRoom, DoubleRoom, SuiteRoom\n'
    '  Booking, MaintenanceRequest, RoomReview, Notification\n\n'
    'LAYER 2 — Services (Business Logic Layer)\n'
    '  RoomManager, StudentManager, BookingManager\n'
    '  MaintenanceManager, NotificationManager, FileManager\n\n'
    'LAYER 3 — Interfaces (Behavioral Contracts)\n'
    '  Bookable, Saveable, Maintainable, Reviewable, Notifiable\n\n'
    'LAYER 4 — Exceptions (Error Handling Layer)\n'
    '  InvalidBookingException, RoomNotAvailableException\n'
    '  InvalidStudentException, DuplicateStudentException\n'
    '  MaintenanceException, UnauthorizedAccessException\n\n'
    'LAYER 5 — GUI (Presentation Layer)\n'
    '  LoginFrame → AdminDashboard / StudentDashboard\n'
    '  RoomPanel, BookingPanel, BookingDialog, ReportsPanel\n'
    '  SplashScreen, UITheme, AlternatingRowRenderer'
)
code_block(arch)
doc.add_page_break()

# ==============================================================================
# 5. FILE STRUCTURE
# ==============================================================================
h1('5. Project File Structure')
structure = (
    'OOP_ Project/\n'
    '|-- src/\n'
    '|   |-- Main.java                        (Entry point)\n'
    '|   |-- models/\n'
    '|   |   |-- AbstractPerson.java          (Abstract base for users)\n'
    '|   |   |-- AbstractRoom.java            (Abstract base for rooms)\n'
    '|   |   |-- Admin.java                   (Admin entity)\n'
    '|   |   |-- Student.java                 (Student entity)\n'
    '|   |   |-- Booking.java                 (Booking record)\n'
    '|   |   |-- SingleRoom.java              (Concrete: 1-person room)\n'
    '|   |   |-- DoubleRoom.java              (Concrete: 2-person room)\n'
    '|   |   |-- SuiteRoom.java               (Concrete: luxury suite)\n'
    '|   |   |-- MaintenanceRequest.java      (Maintenance entity)\n'
    '|   |   |-- Notification.java            (Notification entity)\n'
    '|   |   \-- RoomReview.java              (Review entity)\n'
    '|   |-- services/\n'
    '|   |   |-- RoomManager.java             (Room CRUD + File I/O)\n'
    '|   |   |-- StudentManager.java          (Student CRUD + File I/O)\n'
    '|   |   |-- BookingManager.java          (Booking logic + File I/O)\n'
    '|   |   |-- MaintenanceManager.java      (Maintenance logic)\n'
    '|   |   |-- NotificationManager.java     (Notification logic)\n'
    '|   |   \-- FileManager.java             (Low-level file utilities)\n'
    '|   |-- interfaces/\n'
    '|   |   |-- Bookable.java                (book/cancel/getStatus)\n'
    '|   |   |-- Saveable.java                (save/load)\n'
    '|   |   |-- Maintainable.java            (submitRequest/updateStatus)\n'
    '|   |   |-- Reviewable.java              (addReview/getReviews)\n'
    '|   |   \-- Notifiable.java              (notify/getUnread)\n'
    '|   |-- exceptions/\n'
    '|   |   |-- InvalidBookingException.java\n'
    '|   |   |-- RoomNotAvailableException.java\n'
    '|   |   |-- InvalidStudentException.java\n'
    '|   |   |-- DuplicateStudentException.java\n'
    '|   |   |-- MaintenanceException.java\n'
    '|   |   \-- UnauthorizedAccessException.java\n'
    '|   |-- gui/\n'
    '|   |   |-- UITheme.java                 (Colors, fonts, component factories)\n'
    '|   |   |-- LoginFrame.java              (Login + Registration screen)\n'
    '|   |   |-- AdminDashboard.java          (7-tab admin interface)\n'
    '|   |   |-- StudentDashboard.java        (5-tab student interface)\n'
    '|   |   |-- RoomPanel.java               (Room management tab)\n'
    '|   |   |-- BookingPanel.java            (Admin booking view)\n'
    '|   |   |-- BookingDialog.java           (Student booking popup)\n'
    '|   |   |-- ReportsPanel.java            (Charts and CSV export)\n'
    '|   |   |-- SplashScreen.java            (Startup splash)\n'
    '|   |   \-- AlternatingRowRenderer.java  (Table row colors)\n'
    '|   \-- utils/\n'
    '|       |-- DataStore.java               (Generic container)\n'
    '|       |-- PersonPrinter.java           (Polymorphism demo)\n'
    '|       \-- RoomPrinter.java             (Polymorphism demo)\n'
    '|-- data/                                (Auto-created at runtime)\n'
    '|   |-- rooms.txt\n'
    '|   |-- students.txt\n'
    '|   |-- bookings.txt\n'
    '|   |-- maintenance.txt\n'
    '|   |-- reviews.txt\n'
    '|   \-- notifications.txt\n'
    '|-- out/                                 (Compiled .class files)\n'
    '|-- DATABASE_GUIDE.md\n'
    '|-- PROJECT_GUIDE.md\n'
    '\-- README.md'
)
code_block(structure)
doc.add_page_break()

# ==============================================================================
# 6. KEY SOURCE CODE
# ==============================================================================
h1('6. Key Source Code')
body('The following section presents the most important files that demonstrate core OOP concepts. '
     'Utility and repetitive files are listed in the file structure section above.')

# AbstractPerson
h2('6.1 AbstractPerson.java  —  Abstract Class + Encapsulation (Lab 03, 07)')
code_block(read_src(r'models\AbstractPerson.java'))

# AbstractRoom
h2('6.2 AbstractRoom.java  —  Abstract Class + Polymorphism (Lab 07, 08)')
code_block(read_src(r'models\AbstractRoom.java'))

# Student
h2('6.3 Student.java  —  Inheritance (Lab 06)')
code_block(read_src(r'models\Student.java'))

# Booking
h2('6.4 Booking.java  —  Composition (Lab 05)')
code_block(read_src(r'models\Booking.java'))

# Bookable interface
h2('6.5 Bookable.java  —  Interface (Lab 09)')
code_block(read_src(r'interfaces\Bookable.java'))

# Saveable interface
h2('6.6 Saveable.java  —  Interface with default methods (Lab 09, 12)')
code_block(read_src(r'interfaces\Saveable.java'))

# InvalidBookingException
h2('6.7 InvalidBookingException.java  —  Custom Exception (Lab 10)')
code_block(read_src(r'exceptions\InvalidBookingException.java'))

# Main.java
h2('6.8 Main.java  —  Polymorphism + Generics + File I/O (Lab 08, 11, 12)')
code_block(read_src(r'Main.java'))

doc.add_page_break()

# ==============================================================================
# 7. FEATURES
# ==============================================================================
h1('7. System Features')

h2('7.1 Admin Features')
for f in [
    'Add, edit, and delete room records (Single / Double / Suite)',
    'View all bookings across all students with status indicators',
    'Browse all registered students and their details',
    'Update maintenance request statuses (Pending → In Progress → Resolved)',
    'View and moderate all room reviews submitted by students',
    'Send system-wide notifications to all users',
    'Visual reports dashboard with Java2D bar charts (revenue by room type)',
    'Export booking data to CSV for external use',
    'Live stats bar showing total rooms, bookings, and students',
]:
    bullet(f)

h2('7.2 Student Features')
for f in [
    'Self-registration with ID, name, contact, department, program, semester',
    'Secure login with password masking toggle (show/hide)',
    'Browse available rooms with real-time cost estimation by stay duration',
    'Filter rooms by type: All / Single / Double / Suite',
    'Book a room via interactive dialog with date selection',
    'Cancel active bookings with one click',
    'Submit maintenance requests for room issues',
    'Leave star ratings (1–5) and written reviews for rooms',
    'View personal notifications with unread badge counter',
]:
    bullet(f)

h2('7.3 Common Features')
for f in [
    'Animated splash screen on startup',
    'Gradient blue headers on all dashboards',
    'Hover-effect rounded buttons throughout the UI',
    'Alternating row colors in all data tables',
    'Confirmation dialogs for destructive actions (delete, cancel, logout)',
    'Auto-creation of /data directory and sample data on first run',
]:
    bullet(f)

doc.add_page_break()

# ==============================================================================
# 8. HOW TO RUN
# ==============================================================================
h1('8. How to Compile and Run')

h2('8.1 Requirements')
for r in [
    'Java JDK 11 or later (download from https://adoptium.net)',
    'Windows 10/11 (or any OS with JDK)',
    'No external libraries or database installation required',
]:
    bullet(r)

h2('8.2 Compile and Run (PowerShell)')
code_block(
    'cd "c:\\Users\\Anees\\Desktop\\OOP_ Project"\n\n'
    '# Compile all Java files\n'
    'javac -d out (Get-ChildItem -Path src -Recurse -Filter *.java | ForEach-Object { $_.FullName })\n\n'
    '# Run the application\n'
    'java -cp out Main'
)

h2('8.3 Compile and Run (Command Prompt / bash)')
code_block(
    'cd "OOP_ Project"\n\n'
    '# Compile\n'
    'javac -d out $(find src -name "*.java")\n\n'
    '# Run\n'
    'java -cp out Main'
)

h2('8.4 Default Login Credentials')
ct = doc.add_table(rows=1, cols=3)
ct.style = 'Light Shading Accent 1'
ch = ct.rows[0].cells
for i, v in enumerate(['Role', 'Username / ID', 'Password']):
    ch[i].text = v
    for rn in ch[i].paragraphs[0].runs:
        rn.font.bold = True
for row in [('Admin', 'admin', 'admin123'), ('Student', 'SP23-BSE-030', 'ali12345')]:
    rc = ct.add_row().cells
    for i, v in enumerate(row):
        rc[i].text = v

doc.add_page_break()

# ==============================================================================
# 9. SAMPLE DATA FORMAT
# ==============================================================================
h1('9. Sample Data Format (File I/O — Lab 12)')
body('All data is automatically created in the /data folder on first run. '
     'Each file uses pipe ( | ) as a delimiter.')

h2('9.1 rooms.txt')
code_block('A101|Single|1|8000.0|true|WiFi,Fan\nB201|Double|2|12000.0|true|WiFi,AC,TV\nC301|Suite|3|20000.0|true|WiFi,AC,TV,Kitchen')

h2('9.2 students.txt')
code_block('SP23-BSE-030|Ali Hassan|03001234567|ali12345|CS|BSE|5\nSP23-BSE-031|Sara Khan|03009876543|sara1234|CS|BSE|5')

h2('9.3 bookings.txt')
code_block('BK-001|SP23-BSE-030|A101|2024-01-10|2024-01-15|2024-06-15|Active')

h2('9.4 maintenance.txt')
code_block('MR-001|A101|AC not cooling|Pending|2024-01-12||SP23-BSE-030')

h2('9.5 reviews.txt')
code_block('RV-001|A101|SP23-BSE-030|4|Very comfortable room, clean and quiet.|2024-02-01')

h2('9.6 notifications.txt')
code_block('NF-001|SP23-BSE-030|Your booking BK-001 has been confirmed.|2024-01-10|false')

doc.add_page_break()

# ==============================================================================
# 10. OUTPUT SCREENSHOTS
# ==============================================================================
h1('10. Output Screenshots')
body('The following screenshots show the application running on Windows. '
     'All screens use the custom UITheme with gradient headers, '
     'rounded buttons, and alternating table rows.')

screenshots = [
    ('01_splash.png',        'Figure 1: Splash Screen — Animated progress bar on startup'),
    ('02_login.png',         'Figure 2: Login Screen — Gradient blue sidebar + white card form'),
    ('03_register.png',      'Figure 3: Registration Screen — Student self-registration form'),
    ('04_admin_rooms.png',   'Figure 4: Admin Dashboard — Rooms Tab with room management'),
    ('05_admin_bookings.png','Figure 5: Admin Dashboard — All Bookings Tab'),
    ('06_admin_students.png','Figure 6: Admin Dashboard — Students Tab'),
    ('07_admin_maintenance.png','Figure 7: Admin Dashboard — Maintenance Tab'),
    ('08_admin_reports.png', 'Figure 8: Admin Dashboard — Reports Tab with Java2D bar chart'),
    ('09_student_browse.png','Figure 9: Student Dashboard — Browse Rooms Tab'),
    ('10_booking_dialog.png','Figure 10: Booking Dialog — Date picker and cost calculator'),
    ('11_student_bookings.png','Figure 11: Student Dashboard — My Bookings Tab'),
    ('12_student_maintenance.png','Figure 12: Student Dashboard — Submit Maintenance Tab'),
    ('13_student_reviews.png','Figure 13: Student Dashboard — Reviews Tab with star rating'),
    ('14_notifications.png', 'Figure 14: Student Dashboard — Notifications Tab'),
]

for filename, caption in screenshots:
    h2(caption.split('—')[0].strip() if '—' in caption else caption)
    add_image(filename, caption)
    doc.add_paragraph()

doc.add_page_break()

# ==============================================================================
# 11. FUTURE ENHANCEMENTS
# ==============================================================================
h1('11. Future Enhancements')
for e in [
    'PostgreSQL/MySQL integration via JDBC for concurrent multi-user access.',
    'Firebase Realtime Database for cloud-based data synchronization.',
    'REST API backend to support web and Android frontends.',
    'Email/SMS notification via JavaMail API for booking confirmations.',
    'Payment gateway integration for online room fee collection.',
    'Room photo gallery with image upload support.',
    'Advanced analytics with PDF/Excel export using Apache POI.',
    'Third role: Hostel Warden with partial admin permissions.',
    'Automated room allocation algorithm based on student preferences.',
]:
    bullet(e)

# ==============================================================================
# 12. CONCLUSION
# ==============================================================================
h1('12. Conclusion')
body(
    'The Hostel Room Booking System successfully demonstrates comprehensive application of '
    'all Object-Oriented Programming principles taught in CSC-241. The project implements '
    'encapsulation, inheritance, polymorphism, abstract classes, interfaces, custom exception '
    'handling, generics, Java File I/O, and event-driven GUI programming — all integrated '
    'into a single cohesive application.\n\n'
    'The layered architecture (Models → Services → Interfaces → Exceptions → GUI) ensures '
    'clear separation of concerns, making the codebase maintainable and extensible. '
    'The application runs without any external dependencies, making it fully portable '
    'and easy to demonstrate on any machine with JDK installed.\n\n'
    'This project demonstrates practical software engineering skills including requirement '
    'analysis, object-oriented design, GUI development, data persistence, error handling, '
    'and professional code documentation — all directly applicable to industry-level '
    'Java development.'
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, 'Hostel_System_Project_Report.docx')
doc.save(out)
print('Report saved:', out)
